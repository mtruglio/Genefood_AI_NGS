from posixpath import join
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os
from time import sleep
from shutil import copyfile
from datetime import datetime
import json
from .docx_to_pdf import convert_to, joinpdf, merge_docx
from .filter_JSON import gather_data
from .claude_api_caller import ask_claude
import re
import ast
from .fill_indicazioni_alimentari import fill_template_from_dict

titles_dict = {'Base':'Base', "Plus":"Plus", "Vita":"Vita+", "Ageing":"Food, Aging & Sport", "Sport":"Food & Sport", "Mamma":"Mamma", "Junior_carie":"Junior"}
def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify

    usage:

        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

    provided values are in twentieths of a point (1/1440 of an inch).
    read more here: http://officeopenxml.com/WPtableCellMargins.php
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for m in [
        "top",
        "start",
        "bottom",
        "end",
    ]:
        if m in kwargs:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)

    tcPr.append(tcMar)

def clean_and_convert_to_dict(data):
    # Step 1: Preserve newlines in the "Diagnosi" field
    # Use a regex to find the "Diagnosi" field and temporarily replace newlines with a placeholder
    data = re.sub(r'("Diagnosi"\s*:\s*")(.*?)(")', 
                  lambda m: m.group(1) + m.group(2).replace('\n', '<NEWLINE>') + m.group(3), 
                  data, 
                  flags=re.DOTALL)
    

    # Step 2: Remove all other newlines in the data
    data = data.replace('\n', '')
    # print(data)
    
    # Correct JSON-style true/false/null to Python equivalents
    data = data.replace('true', 'True').replace('false', 'False').replace('null', 'None')

    try:
        # Use ast.literal_eval to safely evaluate the string to a Python dictionary
        python_dict = ast.literal_eval(data)
        return python_dict
    except (SyntaxError, ValueError) as e:
        print(f"Error during conversion: {e}")
        return None


diz_traduttore_pop ={"Ferro Basso":"Carenza di Ferro", "Emocromatosi":'Emocromatosi', 'Low Vitamin B9':"Vitamina B9 (Acido Folico)", 'Low Vitamin D':'Vitamina D',
    "Low Vitamin B12":"Vitamina B12", "Low Vitamin B6":"Vitamina B6", "Low Zinc":"Zinco", "Sodium":"Sodio", 'Sens. Alcol' : "Alcool",'Fruttosio':"Fruttosio", 'Lattosio':'Lattosio','Solfiti':'Solfiti', 
    'Nichel':'Nichel', 'Caffeina':'Caffeina','Glutine':'Glutine', "Crociato Anteriore":"Predisposizione genetica a danni al Legamento Crociato Anteriore", 
    "Crampi-Debolezza Tendinea":"Predisposizione genetica a Crampi/Debolezza Tendinea", "Tendinopatie":"Rischio Tendinopatie", 
    "Sport Resistenza-Potenza":"Predisposizione genetica a Sport Resistenza-Potenza",
    "Danno muscolare":"Rischio Danno muscolare", "Osteoartrosi e fratture":"Rischio Osteoartrosi e fratture", "Infiammazione Cronica":"Aumentata Infiammazione Cronica",
    "Invecchiamento Precoce":"Invecchiamento Precoce", "Memoria Breve":"Memoria a breve termine", 
    "Calo att. antiossidante":"Ridotta attività antiossidante", "Elast. Pelle":"Ridotta elasticità della pelle", 
    "Idrat. Pelle":"Ridotta Idratazione della pelle", "Funzioni cognitive":"Ridotte Funzioni cognitive", 
    "Diabete e ipercolesterolemia":"Rischio Diabete e ipercolesterolemia", "Rischio Cardio":"Rischio patologie cardiovascolari",
    "Calcio Alto":"Eccesso di Calcio", "Carie":"Predisposizione alle carie", "Frag. Ossea":"Rischio Fragilità ossea",
    "Frag. Ossea Pediatrica":"Rischio Fragilità ossea (pediatrica)", "Fosforo Basso":"Carenza di Fosforo", "Sindrome metabolica":"Rischio Sindrome metabolica"
    }

# diz_sport ={
#     "COL27A1-GT":"La presenza dell'allele G predispone alle tendinopatie croniche (Tendine d'Achille)",
#     "TNC-CG": "La presenza dell'allele C predispone alle tendinopatie croniche (Tendine d'Achille)",
#     "TNC-AT": "La presenza dell'allele C predispone alle tendinopatie croniche (Tendine d'Achille)",
#     "COL12A1-CT": "La presenza dell'allele T predispone alla rottura del legamento crociato anteriore",
#     "COL5A1-CT": "La presenza dell'allele T predispone ad una maggiore suscettibilità ai crampi muscolari e debolezza tendinea",
#     "COL1A1-AC": "La presenza del genotipo AA è un fattore di protezione per rottura legamento crociato anteriore",
#     "ESR1-CT": "La presenza dell'allele C fornisce una protezione contro il danno muscolare riducendo la rigidità muscolare",
#     "GDF5-AG": "La presenza dell'allele A è un fattore di rischio per osteoartrosi e fratture ossee",
#     "LRP5-CT": "La presenza dell'allele T è un fattore di rischio per fratture ossee",
#     "MCT1-AT": "La presenza del genotipo AA è associato a un aumentato rischio di lesioni muscolari rispetto agli altri genotipi",
#     "MMP3-CT": "La presenza dell'allele C predispone alle tendinopatie croniche (Tendine d'Achille)",
#     "MMP3-GA": "La presenza dell'allele C predispone alle tendinopatie croniche (Tendine d'Achille)",    
#     "ACTN3-TC": "Influenza la predisposizione agli sport di potenza (CC), intermedi (CT), e di resistenza (CT)",
#     "NOS3-TC": "La presenza dell'allele T porta a una predisposizione agli sport di potenza",
#     "NOS3-GT": "La presenza dell'allele G porta a una predisposizione agli sport di potenza",
#     "TNC-GA": "La presenza dell'allele C predispone alle tendinopatie croniche (Tendine d'Achille)"}

# diz_ageing = {
#     "CRP-rs1205": "Il genotipo CC è associato ad un alto rischio di infiammazione cronica",
#     "IL1B-rs16944": "Il genotipo AA è associato ad una diminizione della memoria a breve termine",
#     "IL1B-rs1143634": "Il genotipo AA è associato funzioni cognitive ridotte",
#     "IL1B-rs1143633": "Il genotipo AA è associato funzioni cognitive ridotte",
#     "IL6-rs1800796": "Il genotipo GG è associato al rischio di osteoartriti",
#     "GPX1-rs1050450": "Il genotipo AA è associato a rischio di malattie cardiache",
#     "SOD3-rs2536512": "Il genotipo GG è associato con una maggiore longevità",
#     "SOD3-rs13306703": "Il genotipo CC è associato a rischio di malattie cardiache",
#     "SOD3-rs699473": "Il genotipo CC è associato a rischio di malattie cardiache",
#     "SOD3-rs1799895": "Il genotipo CC è associato a rischio di malattie cardiache",
#     "CAT-rs1001179": "Il genotipo TT è associato con il rischio di sviluppare forme lievi di diabete e ipercolesterolemia",
#     "NQO1-rs1800566": "Il genotipo AA è associato a rischio di malattie cardiache",
#     "SOD2-rs4880": "Il genotipo AA è associato ad una diminuita attività antiossidante",
#     "SOD2-rs1141718": "Il genotipo AA è associato ad una diminuita attività antiossidante",
#     "AQUAPORIN3-rs17553719": "Il genotipo CC è associato ad una minore idratazione della pelle"
# }


def get_pz_number(patients_dict, code):
    # print(patients_dict)
    for pz in patients_dict:
        if patients_dict[pz]['code']== code:
            return pz
    return "ERROR"

def cambiagenere(testo):
    if testo=='ALTO':
        return 'ALTA'
    if testo == 'MEDIO':
        return 'MEDIA'
    else:
        return testo
        
    
def replace_in_paragraph(document, target, mytext):
    # print("DEBUG replacement", target)
    for paragraph in document.paragraphs:
        # if target == 'test_title':
            # print(paragraph.text)
        if target in paragraph.text:
            # if target == 'test_title':
                # print('found!')
            inline = paragraph.runs
            
            for i in range(len(inline)):
                if target == 'test_title':
                    print(inline[i].text)

                if target in inline[i].text:
                    print('found!')
                    text = inline[i].text.replace(target, mytext)
                    inline[i].text = text

def replace_in_table(document, target, mytext):
    # print("DEBUG replacement", target)

    shading_green = parse_xml(r'<w:shd {} w:fill="008000"/>'.format(nsdecls('w')))
    shading_yellow = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
    shading_orange = parse_xml(r'<w:shd {} w:fill="F49D16"/>'.format(nsdecls('w')))
    shading_red = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))


    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:

                    if target in paragraph.text:
                        inline = paragraph.runs
                        
                        for i in range(len(inline)):
                            # print(inline[i].text)
                            # input()
                            if target in inline[i].text:
                                # print("Found!")
                                text = inline[i].text.replace(target, mytext)
                                inline[i].text = text
                                if "sintesi" not in target:
                                    if mytext == 'NON EVIDENTE':
                                        cell._tc.get_or_add_tcPr().append(shading_green)
                                        run = cell.paragraphs[0].runs[0]
                                        run.font.color.rgb = RGBColor(255,255,255)
                                    elif mytext == 'LIEVE':
                                        cell._tc.get_or_add_tcPr().append(shading_yellow)
                                    elif mytext == 'MEDIO':
                                        cell._tc.get_or_add_tcPr().append(shading_orange)
                                    elif mytext == 'ALTO':
                                        cell._tc.get_or_add_tcPr().append(shading_red)    
                                        run = cell.paragraphs[0].runs[0]
                                        run.font.color.rgb = RGBColor(255,255,255)                                                                        

                        # print("NEW TEXT", paragraph.text)
                    # input()
def paragraph_format_run(cell):
    paragraph = cell.paragraphs[0]
    format = paragraph.paragraph_format
    run = paragraph.add_run()
    
    format.space_before = Pt(0)
    format.space_after = Pt(0)
    format.line_spacing = 1.0
    format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph, format, run

def calculate_age(dob):
    dob = datetime.strptime(dob, "%d/%m/%Y")
    today = datetime.today()
    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
    return age

def assemble_report(analysis_type, patient_id, raw_results, reports, scores_peso, scores_t2d, \
    scores_cardio, scores_mamma, notes_mamma, scores_plus, notes_plus, scores_vita, notes_vita, 
    scores_sport, notes_sport, scores_ageing, notes_ageing, scores_junior_carie, \
    notes_junior_carie, scores_junior_frag, notes_junior_frag, \
    scores_junior_met, notes_junior_met, scores_junior_intoll, \
    notes_junior_intoll, testi, button, debug='off'): 
    
    now = datetime.now()
    print("DEBUG time", now)
    print("TESTI")
    print(testi)
    print(raw_results["Condizioni"])
    print(reports)
    dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
    patient_number = patient_id
    committent= reports[analysis_type][0][patient_number]['committent']     
    complete_json = "static/consolidated_data_italian_with_subcategories.json"
    if analysis_type == "Junior_carie":
        document = Document('static/Referto_Junior.docx')
    else:
        document = Document('static/Referto_{}.docx'.format(analysis_type))

    for i in range(0, len(document.sections)):
        header = document.sections[i].header
        for p in range(0, len(header.paragraphs)):
            htable=header.add_table(1, 2, Inches(7.4))
            htable.alignment = WD_TABLE_ALIGNMENT.LEFT
            htab_cells=htable.rows[0].cells
            ht0=htab_cells[0]
            ht0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell_p, cell_f, cell_r = paragraph_format_run(ht0)

            if committent == 'Genessere':
                cell_r.add_picture("static/header_logos/intestazione_genessere.png", width=Inches(2.55))
                print("Inserted genessere logo")
            elif committent == 'Braincare':
                cell_r.add_picture("static/header_logos/intestazione_braincare.png", width=Inches(2.55))
                print("Inserted braincare logo")
            else:
                if analysis_type == "Junior_carie":
                     cell_r.add_picture("static/header_logos/logo_Base.png".format(analysis_type), width=Inches(1.19))
                else: 
                    cell_r.add_picture("static/header_logos/logo_{}.png".format(analysis_type), width=Inches(1.19))


            ht1=htab_cells[1]
            ht0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell_p, cell_f, cell_r = paragraph_format_run(ht1)
            if committent == 'Longevia':
                cell_r.add_picture("static/header_logos/intestazione_longevia.png", width=Inches(2.55))
            elif committent == 'IkonAcilia':
                cell_r.add_picture("static/header_logos/intestazione_ikonacilia.png", width=Inches(2.55))
            elif committent == 'IkonCasalPalocco':
                cell_r.add_picture("static/header_logos/intestazione_ikoncasalpalocco.png", width=Inches(2.55))
            elif committent == 'IkonFiumicino':
                cell_r.add_picture("static/header_logos/intestazione_ikonfiumicino.png", width=Inches(2.55))
            else:
                cell_r.add_picture("static/header_logos/intestazione_altamedica.png", width=Inches(3.33))
            cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            htable.allow_autofit = True
            ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT




    if debug=='on':
        print(scores_peso, scores_t2d, scores_cardio)

    #ADD TITLE
    if committent == "Genessere":
        replace_in_paragraph(document, 'test_title', 'Genessere') 
    elif committent == "Braincare":
        replace_in_paragraph(document, 'test_title', 'Braincare')
    else:
        replace_in_paragraph(document, 'test_title', 'GENEFOOD {}'.format(titles_dict[analysis_type].upper())) 
        


    # converting df back from dictionary
    for r in raw_results:
        raw_results[r][0] = pd.DataFrame.from_dict(raw_results[r][0])    
    

    if debug=='on':
        print('\n\n\n\n\n\n\n\n\n\n')
        print('### REPORTING PATIENT {0}, analysis type {1} ###'.format(patient_id, analysis_type))
        print('### who is patient {0}, {1} in the dictionary ###'.format(patient_number, reports[analysis_type][0][patient_number]['name']))
            
    name = reports[analysis_type][0][patient_number]['name']
    pz_code = str(patient_id)
    date = dt_string
    dob = reports[analysis_type][0][patient_number]['DOB']
    age = calculate_age(dob)

    if analysis_type=="Mamma":
        gestazione = str(reports[analysis_type][0][patient_number]['gestazione'])
        sex = 'F'
    else:
        sex = reports[analysis_type][0][patient_number]['sesso']
        gestazione = ''
        # Add sig or sig.ra
        if sex == "M":
            replace_in_paragraph(document, 'intestaz_sex', 'Sig.') 
        elif sex == "F":
            replace_in_paragraph(document, 'intestaz_sex', 'Sig.ra') 




    weight = str(reports[analysis_type][0][patient_number]['peso'])
    height = str(reports[analysis_type][0][patient_number]['altezza'])
    
    weight = weight.replace(',','.')
    height = height.replace(',','.')

    bmi = 'NA'
    print("WEIGHT:", weight, type(weight))
    print("HEIGHT:", height, type(height))
    if (weight not in ['0','NA', 'NO']) and (height not in ['0','NA', 'NO']):
        bmi_n = str(round(float(weight)/((float(height)/100)*(float(height)/100)), 2))
        if analysis_type == 'Mamma':
            bmi = bmi_n
        else:
            if float(bmi_n) < 16.5:
                bmi = bmi_n + ' (Sottopeso severo)'
            elif 16.5 <= float(bmi_n) <= 18.4:
                bmi = bmi_n + ' (Sottopeso)'
            elif 18.5 <= float(bmi_n) <= 24.9:
                bmi = bmi_n + ' (Normopeso)'
            elif 25 <= float(bmi_n) <= 30:
                bmi = bmi_n + ' (Sovrappeso)'
            elif 30.1 <= float(bmi_n) <= 34.9:
                bmi = bmi_n + ' (Obesità I grado)'
            elif 35 <= float(bmi_n) <= 40:
                bmi = bmi_n + ' (Obesità II grado)'     
            elif float(bmi_n) > 40:   
                bmi = bmi_n + ' (Obesità III grado)'     

    diete = [] # an empty list of all the diets (pdf) that the program will look for, based on results.
    diete_for_AI = ["Paziente:"]
    base_condition_filter = []
    other_conditions_filter = []
    
    if analysis_type == "Mamma":
        diete_for_AI.append("Donna in gravidanza.")

    else:
        if sex == "M":
            if age < 18 and age >13:
                diete_for_AI.append("Ragazzo.")
            elif age <= 13:
                diete_for_AI.append("Bambino.")
            elif age >= 18:
                diete_for_AI.append("Uomo.")
        elif sex == "F":
            if age < 18 and age >13:
                diete_for_AI.append("Ragazza.")
            elif age <= 13:
                diete_for_AI.append("Bambina.")
            elif age >= 18:
                diete_for_AI.append("Donna.")
    
    diete_for_AI.append("Età: {}.".format(age))
    diete_for_AI.append("BMI: {}.".format(bmi))

    if analysis_type in ['Base', 'Plus', 'Vita', "Sport", "Ageing", "Mamma"]:
        val_peso=reports['Base'][2][patient_number]['Peso']
        val_t2d=reports['Base'][2][patient_number]['T2D']
        val_cardio = reports['Base'][2][patient_number]['Cardio']
        
        # Adding dieta base (from T2D, Cardio and Peso scores we already know)
        diete.append("./static/indicazioni_alimentari/T2D_{0}+Cardio_{1}+Peso_{2}.pdf".format(val_t2d.title().replace(' ',''), val_cardio.title().replace(' ',''), val_peso.title().replace(' ','')))
        base_condition_filter.append(f"T2D_{val_t2d.title().replace(' ','')}")
        base_condition_filter.append(f"Cardio_{val_cardio.title().replace(' ','')}")
        base_condition_filter.append(f"Peso_{val_peso.title().replace(' ','')}")
        string_baseconditions = ''
        if val_t2d.title().replace(' ','') != 'NonEvidente':
            string_baseconditions += f"Predisposizione al Diabete di tipo 2 ({val_t2d.title().replace(' ','')}), "
        else:
            string_baseconditions += "Non evidente predisposizione al Diabete di tipo 2, "
        if val_cardio.title().replace(' ','') != 'NonEvidente':
            string_baseconditions += f"Predisposizione alle malattie cardiovascolari ({val_cardio.title().replace(' ','')}), "
        else:
            string_baseconditions += "Non evidente predisposizione alle malattie cardiovascolari, "
        if val_peso.title().replace(' ','') != 'NonEvidente':
            string_baseconditions += f"Predisposizione all'aumento di peso ({val_peso.title().replace(' ','')})."
        else:
            string_baseconditions += "Non evidente predisposizione all'aumento di peso."
        diete_for_AI.append(string_baseconditions)

    if analysis_type in ['Plus', 'Vita', "Sport", "Ageing", "Mamma", "Junior_carie"]:
        testo_intolleranzeshort = testi[0]
        testo_intolleranzelong = testi[1]
        testo_assorbimentoshort = testi[2]
        testo_assorbimentolong = testi[3]
        testo_sportshort = testi[4]
        testo_sportlong = testi[5]
        testo_ageingshort = testi[6]
        testo_ageinglong = testi[7]
        testo_juniormet_short = testi[8]
        testo_juniormet_long = testi[9]
        testo_juniorintoll_short = testi[10]
        testo_juniorintoll_long = testi[11]
        testo_juniorfrag_short = testi[12]
        testo_juniorfrag_long = testi[13]
        testo_juniorcarie_short = testi[14]
        testo_juniorcarie_long = testi[15]
        print("Testi in ricerca",testo_juniorcarie_long, testo_juniorintoll_long, testo_juniorfrag_long, testo_juniormet_long)
        # Predisposizioni
        for v in (['<testo_intolleranze>', testo_intolleranzeshort], ['<testo_assorbimento>',testo_assorbimentoshort], \
            ['<testo_sport>',testo_sportshort], ['<testo_ageing>',testo_ageingshort], ['<testo_sindrome_met>', testo_juniormet_short], \
            ['<testo_intolleranze_junior>', testo_juniorintoll_short], ['<testo_fragilita>', testo_juniorfrag_short], \
            ['<testo_carie>', testo_juniorcarie_short]):
            replace_in_table(document, v[0], v[1])

        for v in (['testo_intolleranze_lungo', testo_intolleranzelong], ['testo_assorbimento_lungo',testo_assorbimentolong], \
            ['testo_sport_lungo', testo_sportlong], ['testo_ageing_lungo',testo_ageinglong], ['testo_sindrome_met_lungo', testo_juniormet_long], \
            ['testo_intolleranze_junior_lungo', testo_juniorintoll_long], ['testo_fragilita_lungo', testo_juniorfrag_long], \
            ['testo_carie_lungo', testo_juniorcarie_long]):
            replace_in_paragraph(document, v[0], v[1])  

        if testo_intolleranzeshort!="" and testo_intolleranzeshort!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_intolleranzeshort.lower().capitalize()+".")
        if "INTOLLERANZA AL GLUTINE" not in testo_intolleranzeshort and not pd.isna(raw_results["Condizioni"][0][patient_number][0]) and "glutine" not in raw_results["Condizioni"][0][patient_number][0].lower() and "celiachia" not in raw_results["Condizioni"][0][patient_number][0].lower():
            diete_for_AI.append("No intolleranza al Glutine.")    
        if "INTOLLERANZA AL LATTOSIO" not in testo_intolleranzeshort and not pd.isna(raw_results["Condizioni"][0][patient_number][0]) and "lattosio" not in raw_results["Condizioni"][0][patient_number][0].lower():
            diete_for_AI.append("No intolleranza al Lattosio.")

        if testo_assorbimentoshort!="" and testo_assorbimentoshort!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_assorbimentoshort.lower().capitalize()+".")

        if testo_sportshort!="" and testo_sportshort!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_sportshort.lower().capitalize()+".")
        
        if testo_ageingshort!="" and testo_ageingshort!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_ageingshort.lower().capitalize()+".")
            
        if testo_juniormet_short!="" and testo_juniormet_short!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_juniormet_short.lower().capitalize()+".")
        if testo_juniorintoll_short!="" and testo_juniorintoll_short!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_juniorintoll_short.lower().capitalize()+".")
        if testo_juniorfrag_short!="" and testo_juniorfrag_short!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_juniorfrag_short.lower().capitalize()+".")
        if testo_juniorcarie_short!="" and testo_juniorcarie_short!="NON SI EVIDENZIANO PREDISPOSIZIONI GENETICHE":
            diete_for_AI.append(testo_juniorcarie_short.lower().capitalize()+".")
        
            
    # Nome paziente
    replace_in_paragraph(document, '<nome>', name)
      


    # Anagrafica
    for v in (['<NrAcc>', pz_code], ['<timestamp>', date],['<datanasc>', dob], ['<sex>', sex], ['<gestazione>', gestazione], ['<weight>', weight], ['<height>', height], ['<bmi>', bmi]):
        print("DEBUG", v)
        replace_in_table(document, v[0], v[1])



    # For each format, the cardio-T2D-peso tables have a different table number.
    # This happens because the summary tables in the first page are variable
    
    if analysis_type == 'Base':
        first_gen_table = 5
    elif analysis_type == 'Plus':
        first_gen_table = 6
    elif analysis_type == 'Vita' or analysis_type == 'Mamma':
        first_gen_table = 7
    elif analysis_type == 'Sport':
        first_gen_table = 8
    elif analysis_type == 'Ageing':
        first_gen_table = 9
    elif analysis_type == 'Junior_carie':
        first_gen_table = 5

    if analysis_type!='Junior_carie':
        # Summary one-line tables for peso t2d and cardio (first page).
        for v in (['<valpeso_sintesi>', val_peso], ['<valt2d_sintesi>',val_t2d],['<valcardio_sintesi>', val_cardio]):
            replace_in_table(document, v[0], cambiagenere(v[1].upper()))

        # Risk summary above each gene table for peso, t2d, cardio (following pages)    
        for v in (['<valpeso>', val_peso], ['<valt2d>',val_t2d],['<valcardio>', val_cardio]):
            replace_in_table(document, v[0], v[1].upper()) 

    count=0
    for table in document.tables:
        if debug=='on':
            print('starting table {}'.format(count))
        
        #Filling Peso gene table
        if count==first_gen_table:
            if analysis_type!='Junior_carie':
                if debug=='on':
                    print("Tabella Base - peso")
                for i,row in raw_results['Base'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_peso and row[patient_number].lower()!='no':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']
                        if analysis_type!='Mamma': # Mamma is all shaded in dark pink for some reason.
                            shading = parse_xml(r'<w:shd {} w:fill="F49D16"/>'.format(nsdecls('w')))
                        else:
                            shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))

                        cells[0]._tc.get_or_add_tcPr().append(shading)

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True
                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
            else:
                if debug=='on':
                    print("Tabella Junior met")
                for i,row in raw_results['Junior_sindrome_met'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_junior_met and row[patient_number].lower()!='no' and row['Gene']!='MTHFR':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']
                        
                        shading = parse_xml(r'<w:shd {} w:fill="CC2A58"/>'.format(nsdecls('w')))


                        cells[0]._tc.get_or_add_tcPr().append(shading)

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True
                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255) 

                for param in reports['Junior_sindrome_met'][2][patient_number]:
                    param_result_value = reports['Junior_sindrome_met'][2][patient_number][param]
                    if param_result_value != False:
                        base_condition_filter.append(f"Junior+{param.replace(' ','')}_{param_result_value}")                                                                                      

        if count==first_gen_table+1:
            if analysis_type=='Junior_carie':
                if debug=='on':
                    print("Tabella Junior met, MTHFR")
                for i,row in raw_results['Junior_sindrome_met'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_junior_met and row[patient_number].lower()!='no' and row['Gene']=='MTHFR':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']
                        
                        shading = parse_xml(r'<w:shd {} w:fill="CC2A58"/>'.format(nsdecls('w')))


                        cells[0]._tc.get_or_add_tcPr().append(shading)


                        cells[1].text = row['SNP']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
                        
                        cells[2].text = row['WT']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        cells[3].text = row['alt']
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        cells[4].text = row[patient_number]
                        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[4].paragraphs[0].runs[0]
                        run.font.bold = True
                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)  
    
        if count==first_gen_table+2:
            if analysis_type!='Junior_carie':
                if debug=='on':
                    print("Tabella Base - T2D")
                
                #Filling T2D gene table
                for i,row in raw_results['Base'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_t2d and row[patient_number].lower()!='no':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']
                        
                        if analysis_type!='Mamma':
                            shading = parse_xml(r'<w:shd {} w:fill="F49D16"/>'.format(nsdecls('w')))
                        else:
                            shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))                    
                            
                        cells[0]._tc.get_or_add_tcPr().append(shading) 

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True

                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255) 
            else:
                if debug=='on':
                    print("Tabella Intol Junior")
                for i,row in raw_results['Junior_intolleranze'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_junior_intoll and row[patient_number].lower()!='no' and notes_junior_intoll[row['Gene']][row['SNP']][0]!='Caffeina':
                        cells = table.add_row().cells
                        
                        cells[0].text = row['Gene']

                        shading = parse_xml(r'<w:shd {} w:fill="92CF50"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(shading) 

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True

                        celltext =[]
                        for note in notes_junior_intoll[row['Gene']][row['SNP']]:
                            celltext.append(diz_traduttore_pop[note])
                        cells[4].text = ', '.join(celltext)
                        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cells[4].paragraphs[0].runs[0]
                        run.font.name = 'Arial'
                        
                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)                     
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER   

                for param in reports['Junior_intolleranze'][2][patient_number]:
                    param_result_value = reports['Junior_intolleranze'][2][patient_number][param]
                    if param_result_value != 'No' and param_result_value != False:
                        other_conditions_filter.append(f"Junior+{param.replace(' ','')}_{param_result_value}") 

        if count==first_gen_table+3 and analysis_type=='Junior_carie':
            if debug=='on':
                print("Tabella Intol Junior, Caffeina")
            for i,row in raw_results['Junior_intolleranze'][0].iterrows():
                # For each row
                if row['Gene'] in scores_junior_intoll and row[patient_number].lower()!='no' and notes_junior_intoll[row['Gene']][row['SNP']][0]=='Caffeina':
                    cells = table.add_row().cells
                    
                    cells[0].text = row['Gene']

                    shading = parse_xml(r'<w:shd {} w:fill="92CF50"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading) 

                    cells[1].text = row['WT']
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[2].text = row['alt']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row[patient_number]
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[3].paragraphs[0].runs[0]
                    run.font.bold = True

                    celltext =[]
                    for note in notes_junior_intoll[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[4].text = ', '.join(celltext)
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cells[4].paragraphs[0].runs[0]
                    run.font.name = 'Arial'
                    
                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)                     
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER             

                        
        if count==first_gen_table+4:
            if analysis_type!='Junior_carie':
                if debug=='on':
                    print("Tabella Base - Cardio")
                
                #Filling Cardio gene table
                for i,row in raw_results['Base'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_cardio and row[patient_number].lower()!='no' and row['Gene']!='MTHFR':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']

                        if analysis_type!='Mamma':
                            shading = parse_xml(r'<w:shd {} w:fill="F49D16"/>'.format(nsdecls('w')))
                        else:
                            shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))

                        cells[0]._tc.get_or_add_tcPr().append(shading)

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True

                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
            else:
                if debug=='on':
                    print("Tabella Junior Carie")
                for i,row in raw_results['Junior_carie'][0].iterrows():
                    if row['Gene'] in scores_junior_carie and row[patient_number].lower()!='no':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']
                        
                        shading = parse_xml(r'<w:shd {} w:fill="82CBD9"/>'.format(nsdecls('w')))


                        cells[0]._tc.get_or_add_tcPr().append(shading)

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True
                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)                    
                for param in reports['Junior_carie'][2][patient_number]:
                    param_result_value = reports['Junior_carie'][2][patient_number][param]
                    if param_result_value != 'No' and param_result_value != False:
                        other_conditions_filter.append(f"Junior+{param.replace(' ','')}_{param_result_value}") 

        if count==first_gen_table+5:
            if analysis_type!='Junior_carie':
                for i,row in raw_results['Base'][0].iterrows():
                    # For each row
                    if row['Gene'] in scores_cardio and row[patient_number].lower()!='no' and row['Gene']=='MTHFR':
                        cells = table.add_row().cells
                        
                        cells[0].text = row['Gene']
                        if analysis_type!='Mamma':
                            shading = parse_xml(r'<w:shd {} w:fill="F49D16"/>'.format(nsdecls('w')))
                        else:
                            shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(shading) 

                        
                        cells[1].text = row['SNP']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            
                        cells[2].text = row['WT']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row['alt']
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[4].text = row[patient_number]
                        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        run = cells[4].paragraphs[0].runs[0]
                        run.font.bold = True

                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)
            else:
                if debug=='on':
                    print("Tabella Junior Fragilita")
                for i,row in raw_results['Junior_fragilita'][0].iterrows():
                    if row['Gene'] in scores_junior_frag and row[patient_number].lower()!='no':
                        cells = table.add_row().cells
                        
                        
                        cells[0].text = row['Gene']
                        shading_blue = parse_xml(r'<w:shd {} w:fill="3BA1DA"/>'.format(nsdecls('w')))
                        cells[0]._tc.get_or_add_tcPr().append(shading_blue) 

                        # celltext =[]
                        # for note in notes_junior_frag[row['Gene']][row['SNP']]:
                        #     celltext.append(diz_traduttore_pop[note])
                        # cells[1].text = ', '.join(celltext)
                        # run = cells[1].paragraphs[0].runs[0]
                        # run.font.size = Pt(9)
                        # run.font.name = 'Arial'
                        # set_cell_margins(cells[1],  top=50, start=50, bottom=50, end=50)
                        
                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True

                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255)  

                for param in reports['Junior_fragilita'][2][patient_number]:
                    param_result_value = reports['Junior_fragilita'][2][patient_number][param]
                    if param_result_value != 'No' and param_result_value != False:
                        other_conditions_filter.append(f"Junior+{param.replace(' ','')}_{param_result_value}") 
                                 
        
        if analysis_type in ['Plus', 'Vita', 'Sport', 'Ageing', 'Mamma'] and count==first_gen_table+6:
            if debug=='on':
                print("Tabella Intolleranze")
            for i,row in raw_results['Plus'][0].iterrows():
                # For each row
                if row['Gene'] in scores_plus and row[patient_number].lower()!='no' and notes_plus[row['Gene']][row['SNP']][0]!='Caffeina':
                    cells = table.add_row().cells
                    
                    cells[0].text = row['Gene']
                    if analysis_type!='Mamma':
                        shading = parse_xml(r'<w:shd {} w:fill="92CF50"/>'.format(nsdecls('w')))
                    else:
                        shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading) 

                    cells[1].text = row['WT']
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[2].text = row['alt']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row[patient_number]
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[3].paragraphs[0].runs[0]
                    run.font.bold = True

                    celltext =[]
                    for note in notes_plus[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[4].text = ', '.join(celltext)
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cells[4].paragraphs[0].runs[0]
                    run.font.name = 'Arial'
                    
                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)                     
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for param in reports['Plus'][2][patient_number]:
                param_result_value = reports['Plus'][2][patient_number][param]
                if param_result_value != 'No' and param_result_value != False and 'Vedi' not in str(param_result_value):
                    if param == "Glutine" or param == "Lattosio":
                        # This is a special case: if Glutine==True or Lattosio==True, diet can be 
                        # 'singola' if cardio==False + Peso==False + T2D==False
                        # 'in abbinamento' if one or more of the three above is True.
                        if debug=='on':
                            print("{}!".format(param), val_peso, val_cardio, val_t2d)
                        if val_peso=='non evidente' and val_cardio=='non evidente' and val_t2d=='non evidente':
                            diete.append('./static/indicazioni_alimentari/{}_singola.pdf'.format(param))
                        else:
                            diete.append('./static/indicazioni_alimentari/{}_in_abbinamento.pdf'.format(param))
                        
                    else:                            
                        diete.append('./static/indicazioni_alimentari/{0}_{1}.pdf'.format(param.replace(' ',''), param_result_value))
                        other_conditions_filter.append(f"{param.replace(' ','')}_{param_result_value}")

        if analysis_type in ['Plus', 'Vita', 'Sport', 'Ageing', 'Mamma'] and count==first_gen_table+7:
            if debug=='on':
                print("Tabella Caffeina")
            for i,row in raw_results['Plus'][0].iterrows():
                # For each row
                if row['Gene'] in scores_plus and row[patient_number].lower()!='no' and notes_plus[row['Gene']][row['SNP']][0]=='Caffeina':
                    cells = table.add_row().cells
                    
                    cells[0].text = row['Gene']
                    if analysis_type!='Mamma':
                        shading = parse_xml(r'<w:shd {} w:fill="92CF50"/>'.format(nsdecls('w')))
                    else:
                        shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading) 

                    cells[1].text = row['WT']
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[2].text = row['alt']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row[patient_number]
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[3].paragraphs[0].runs[0]
                    run.font.bold = True

                    celltext =[]
                    for note in notes_plus[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[4].text = ', '.join(celltext)
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cells[4].paragraphs[0].runs[0]
                    run.font.name = 'Arial'
                    
                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)                     
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 

        if analysis_type in ['Vita', 'Sport', 'Ageing', 'Mamma'] and count==first_gen_table+8:
            if debug=='on':
                print("Tabella Assorbimento")
                print(scores_vita)
            for i,row in raw_results['Vita'][0].iterrows():
                # For each row
                if debug=='on':
                    print("working with", row['Gene'])
                if row['Gene'] in scores_vita and row[patient_number].lower()!='no' and row['Gene']!='MTHFR':
                    cells = table.add_row().cells
                    cells[0].text = row['Gene']

                    if analysis_type!='Mamma':
                        shading = parse_xml(r'<w:shd {} w:fill="CC2A58"/>'.format(nsdecls('w')))
                    else:
                        shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))

                    cells[0]._tc.get_or_add_tcPr().append(shading) 

                    cells[1].text = row['WT']
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[2].text = row['alt']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row[patient_number]
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[3].paragraphs[0].runs[0]
                    run.font.bold = True

                    celltext =[]
                    for note in notes_vita[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[4].text = ', '.join(celltext)
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cells[4].paragraphs[0].runs[0]
                    run.font.name = 'Arial'

                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255) 

            for param in reports['Vita'][2][patient_number]:
                param_result_value = reports['Vita'][2][patient_number][param]
                if param_result_value != 'No' and param_result_value != False and 'Vedi' not in str(param_result_value):
                    diete.append('./static/indicazioni_alimentari/{0}_{1}.pdf'.format(param.replace(' ',''), param_result_value))
                    other_conditions_filter.append(f"{param.replace(' ','')}_{param_result_value}")

            if analysis_type=="Mamma" and row[patient_number].lower()!='no':
                for i,row in raw_results['Mamma'][0].iterrows():
                # For each row
                    if row['Gene'] in scores_mamma and row[patient_number].lower()!='no':
                        cells = table.add_row().cells
                        cells[0].text = row['Gene']

                        shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))

                        cells[0]._tc.get_or_add_tcPr().append(shading) 

                        cells[1].text = row['WT']
                        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[2].text = row['alt']
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        cells[3].text = row[patient_number]
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                        run = cells[3].paragraphs[0].runs[0]
                        run.font.bold = True

                        celltext =[]
                        for note in notes_mamma[row['Gene']][row['SNP']]:
                            celltext.append(diz_traduttore_pop[note])
                        cells[4].text = ', '.join(celltext)                        
                        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = cells[4].paragraphs[0].runs[0]
                        run.font.name = 'Arial'

                        run = cells[0].paragraphs[0].runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255,255,255) 

                for param in reports['Mamma'][2][patient_number]:
                    param_result_value = reports['Mamma'][2][patient_number][param]
                    if param_result_value != 'No' and param_result_value != False and 'Vedi' not in str(param_result_value):
                        diete.append('./static/indicazioni_alimentari/{0}_{1}.pdf'.format(param.replace(' ',''), param_result_value))
                        other_conditions_filter.append(f"{param.replace(' ','')}_{param_result_value}")
                
                # Consigli Gravidanza will always be appended at the end, in any case.
                diete.append('./static/indicazioni_alimentari/Consigli_Nutrizionali_Gravidanza.pdf')
                other_conditions_filter.append("Gravidanza")


        if analysis_type in ['Vita', 'Sport', 'Ageing', 'Mamma'] and count==first_gen_table+9:
            if debug=='on':
                print("Tabella MTHFR vita")
                print(scores_vita)
            for i,row in raw_results['Vita'][0].iterrows():
                # For each row
                if debug=='on':
                    print("working with", row['Gene'])
                if row['Gene'] in scores_vita and row[patient_number].lower()!='no' and row['Gene']=='MTHFR':
                    cells = table.add_row().cells
                    
                    cells[0].text = row['Gene']
                    if analysis_type!='Mamma':
                        shading = parse_xml(r'<w:shd {} w:fill="CC2A58"/>'.format(nsdecls('w')))
                    else:
                        shading = parse_xml(r'<w:shd {} w:fill="C39291"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading) 

                    
                    cells[1].text = row['SNP']
                    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        
                    cells[2].text = row['WT']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row['alt']
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[4].text = row[patient_number]
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    run = cells[4].paragraphs[0].runs[0]
                    run.font.bold = True

                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)                    
                    

        if analysis_type in ['Sport', 'Ageing'] and count==first_gen_table+10:
            if debug=='on':
                print("Tabella Sport")
                print(scores_sport)

            for i,row in raw_results['Sport'][0].iterrows():
                # For each row
                if debug=='on':
                    print("working with", row['Gene'])
                if row['Gene'] in scores_sport and row[patient_number].lower()!='no':
                    cells = table.add_row().cells
                    cells[0].text = row['Gene']
                    shading_azure = parse_xml(r'<w:shd {} w:fill="82CBD9"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading_azure) 

                    celltext =[]
                    for note in notes_sport[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[1].text = ', '.join(celltext)

                    run = cells[1].paragraphs[0].runs[0]
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
                    set_cell_margins(cells[1],  top=50, start=50, bottom=50, end=50)
                    
                    cells[2].text = row['WT']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row['alt']
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[4].text = row[patient_number]
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[4].paragraphs[0].runs[0]
                    run.font.bold = True

                    genekey = row['Gene']+'-'+row['WT']+row['alt']
                    if debug=='on':
                        print("Looking for ", row['Gene'], row['SNP'])


                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)   

            for param in reports['Sport'][2][patient_number]:
                param_result_value = reports['Sport'][2][patient_number][param]
                if param_result_value != 'No' and param_result_value != False and 'Vedi' not in str(param_result_value):
                    diete.append('./static/indicazioni_alimentari/{0}_{1}.pdf'.format(param.replace(' ',''), param_result_value))
                    other_conditions_filter.append(f"{param.replace(' ','')}_{param_result_value}")

        if analysis_type in ['Ageing'] and count==first_gen_table+11:
            if debug=='on':
                print("Tabella Ageing")

            for i,row in raw_results['Ageing'][0].iterrows():
                # For each row
                if row['Gene'] in scores_ageing and row[patient_number].lower()!='no':
                    cells = table.add_row().cells
                    
                    
                    cells[0].text = row['Gene']
                    shading_blue = parse_xml(r'<w:shd {} w:fill="3BA1DA"/>'.format(nsdecls('w')))
                    cells[0]._tc.get_or_add_tcPr().append(shading_blue) 

                    celltext =[]
                    for note in notes_ageing[row['Gene']][row['SNP']]:
                        celltext.append(diz_traduttore_pop[note])
                    cells[1].text = ', '.join(celltext)
                    run = cells[1].paragraphs[0].runs[0]
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
                    set_cell_margins(cells[1],  top=50, start=50, bottom=50, end=50)
                    
                    cells[2].text = row['WT']
                    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[3].text = row['alt']
                    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cells[4].text = row[patient_number]
                    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = cells[4].paragraphs[0].runs[0]
                    run.font.bold = True

                    run = cells[0].paragraphs[0].runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255,255,255)   

            for param in reports['Ageing'][2][patient_number]:
                param_result_value = reports['Ageing'][2][patient_number][param]
                if param_result_value != 'No' and param_result_value != False and 'Vedi' not in str(param_result_value):
                    diete.append('./static/indicazioni_alimentari/{0}_{1}.pdf'.format(param.replace(' ',''), param_result_value))
                    other_conditions_filter.append(f"{param.replace(' ','')}_{param_result_value}")  

        count+=1

                    
    # input("inizio Footer")
    for section in document.sections:
        # section = document.sections[0]
        footer = section.footer
        # for p in footer.paragraphs:
            # if debug=='on':
            #     print(p.text)
        # input("qua")
        
        for t in footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if 'footernote' in paragraph.text:
                            inline = paragraph.runs
                            
                            for i in range(len(inline)):
                                # if debug=='on':
                                #     print(inline[i].text)
                                if 'footernote' in inline[i].text:
                                    if committent == 'Longevia':
                                        text = inline[i].text.replace('footernote', "Paziente: {0} - Referto emesso il {1}".format(name, date))
                                    else:
                                        text = inline[i].text.replace('footernote', "Paziente: {0} - Referto emesso il {1} dal Centro Altamedica di Roma".format(name, date))
                                    inline[i].text = text


    docx_name = './ARCHIVIO/{0}_{1}_{2}_genetics.docx'.format(name, patient_id, analysis_type)
    document.save(docx_name)
    print("DOCX SAVED")
    convert_to('./ARCHIVIO/', docx_name )
    print("CONVERTED TTO PDF (NO DIET)")
    original_pdf = docx_name.replace('.docx','.pdf')
    
    # Add vegan or vegetarian to each pdf name if needed
    if reports[analysis_type][0][patient_number]['lim'] in ['vegano', 'vegana', 'vegetariano', 'vegetariana']:
        for i in range(0, len(diete)):
            diete[i] = diete[i].replace('.pdf', '+'+reports[analysis_type][0][patient_number]['lim'][:-1].capitalize()+'.pdf')
    if reports[analysis_type][0][patient_number]['lim'].lower()!="no":
        diete_for_AI.append(reports[analysis_type][0][patient_number]['lim'][:-1].capitalize()+'.')
    diete.insert(0, original_pdf) # prepend the original results pdf at the beginning of pdf reports list
    if not pd.isna(raw_results["Condizioni"][0][patient_number][0]):
        diete_for_AI.append(raw_results["Condizioni"][0][patient_number][0]+".")
    print("######### Report: ", diete)
    final_pdf = original_pdf.replace('_result.pdf', '.pdf')

    prompt = ' '.join(diete_for_AI)

    # Redundant check for gluten and lactose intolerance. If the previous check was not enough, we check again for absence of any mention of gluten or lactose.
    if "glutine" not in prompt.lower() and "celiachia" not in prompt.lower():
            diete_for_AI.append("No intolleranza al Glutine. ")
    if "lattosio" not in prompt.lower():
            diete_for_AI.append("No intolleranza al Lattosio. ")
    
    template_indicazioni = 'static/Indicazioni_alimentari.docx'
    if "No intolleranza al Glutine" in prompt and "No intolleranza al Lattosio" in prompt:
        template_indicazioni = 'static/Indicazioni_alimentari.docx'
    elif "No intolleranza al Glutine" not in prompt and "No intolleranza al Lattosio" in prompt and ("glutine" in prompt.lower() or "celiachia" in prompt.lower() or "celiac" in prompt.lower()):
        template_indicazioni = 'static/Indicazioni_alimentari_glut.docx'
    elif "No intolleranza al Glutine" in prompt and "No intolleranza al Lattosio" not in prompt and ("lattosio" in prompt.lower()):
        template_indicazioni = 'static/Indicazioni_alimentari_latt.docx'
    elif "No intolleranza al Glutine" not in prompt and "No intolleranza al Lattosio" not in prompt and ("glutine" in prompt.lower() or "celiachia" in prompt.lower() or "celiac" in prompt.lower()) and ("lattosio" in prompt.lower()):
        template_indicazioni = 'static/Indicazioni_alimentari_glut_latt.docx'
    
    print("TEMPLATE: ", template_indicazioni)


        
    diete_for_AI.append("\n1. Dammi la risposta in formato JSON, senza altro testo prima o dopo. Sanifica tutti i caratteri speciali, specialmente le virgolette.\n2. La struttura del JSON in output ricalchera' quella dei JSON di esempio.\n3. Per i pazienti VEGETARIANI, la chiave \"Proteine\" in \"raccomandazioni\" non conterra' le chiavi \"CARNE\" e \"PESCE\", ma solo \"ALIMENTI DI ORIGINE VEGETALE\", \"LEGUMI\", \"LATTICINI\" e \"UOVA\". Per i pazienti VEGANI, la sezione \"Proteine\" in \"raccomandazioni\" non conterra' le chiavi \"CARNE\", \"PESCE\", \"LATTICINI\" e \"UOVA\", ma solo \"ALIMENTI DI ORIGINE VEGETALE\" e \"LEGUMI\".\n4. Nella chiave \"raccomandazioni\" aggiungi eventuali integratori (SENZA SPECIFICARE LA MARCA COMMERCIALE), dosaggio e motivazione, sotto la chiave \"Integratori\", se necessari al paziente nelle sue condizioni. \"Integratori\" sara' strutturato nel seguente modo: \"Integratori\":[{ \"tipo\", \"dosaggio\", \"motivazione\" }].\n4.1 NON PRESCRIVERE INTEGRATORI A BASE DI CROMO. \n4.2 Se prescrivi integratori a base di Omega-3, nel dosaggio inserisci la nota \"La FDA consiglia di non superare un'assunzione totale di 3 grammi al giorno di EPA e DHA combinati (acidi grassi omega-3), con un massimo di 2 grammi provenienti da integratori alimentari.\".\n5. Per la sezione \"diagnosi\", approfondisci in maniera divulgativa ma estensivamente tutti gli aspetti e le eventuali predisposizioni/intolleranze, facendo cenni anche alla genetica, ricalcando i JSON di esempio quando disponibili per una particolare condizione, e solo partendo da questo testo ricalcato eventualmente espanderlo con osservazioni scientificamente fondate. Se nel JSON di esempio sono presenti \"raccomandazioni\" relative a una certa intolleranza/condizione, riportale, e poi espandile se necessario o se pensi di poter aggiungere informazioni importanti.\n6. Se consigli dispositivi medici o trattamenti terapeutici, inseriscili SEMPRE usando la dicitura \"Oltre alla dieta e all'esercizio, potrebbe essere necessaria la prescrizione da parte di un medico specialista di {dispositivo o trattamento suggerito}\" \n7. Nomina i geni solo se gia' presenti nel file JSON di esempio per quella condizione, non aggiungere altri geni non verificati nella tua risposta. Non inserire KCNJ11 tra i geni per il diabete, perche' al momento non lo indaghiamo. Tra i geni per le malattie cardiovascolari, nomina solo APOE e MTNR1B. \n8. Non menzionare cancro o tumore nella sezione \"Diagnosi\".\n9. Dopo aver generato il JSON, validane internamente la struttura prima di fornirmelo.\n10. Espandi le liste alimenti il piu' possibile, cerca di indicare almeno 5-10 alimenti per categoria di macronutrienti e livello di tollerabilita'.\n11. Dopo aver generato le liste di alimenti, confrontale di nuovo con le \"condizioni\" del paziente e correggi eventuali incoerenze.\n12. Hai oltre 8000 token disponibili, elabora estensivamente ogni condizione nella sezione Diagnosi.")

    prompt = ' '.join(diete_for_AI)
    print(raw_results["Condizioni"][0][patient_number][0].lower())
    if "vega" in raw_results["Condizioni"][1][patient_number]['lim'].lower() or "vega" in raw_results["Condizioni"][0][patient_number][0].lower():
        base_condition_filter.append("Vegano")
    if "vege" in raw_results["Condizioni"][1][patient_number]['lim'].lower() or "vegetari" in raw_results["Condizioni"][0][patient_number][0].lower():
        base_condition_filter.append("Vegetariano")    
    print(prompt)
    print(base_condition_filter)
    print(other_conditions_filter)

    gather_data('./static/consolidated_data_italian_with_subcategories.json', './static/consolidated_data_italian_with_subcategories_special.json', base_condition_filter, other_conditions_filter,'./ARCHIVIO/{0}_{1}_{2}_reduced.json'.format(name, patient_id, analysis_type))
    # input("STOP")
    ai_response = ask_claude('./ARCHIVIO/{0}_{1}_{2}_reduced.json'.format(name, patient_id, analysis_type), prompt, analysis_type)
    ai_response_text = ai_response[0].text if isinstance(ai_response, list) and ai_response else ai_response

    # # Write the text to a file # ONLY FOR TESTING
    # with open('./static/{0}_{1}_{2}_ai_response_text.txt'.format(name, patient_id, analysis_type), 'w', encoding='utf-8') as f:
    #     f.write(str(ai_response_text))

    ai_response_dict = clean_and_convert_to_dict(ai_response_text)
    
    # Write the dictionary to a file # ONLY FOR TESTING
    # with open('./static/{0}_{1}_{2}_ai_response_dict.txt'.format(name, patient_id, analysis_type), 'w', encoding='utf-8') as f:
    #     f.write(str(ai_response_dict))

    # Read the dictionary back from the file # ONLY FOR TESTING
    # with open('./static/{0}_{1}_{2}_ai_response_dict.txt'.format(name, patient_id, analysis_type), 'r', encoding='utf-8') as f:
    #     ai_response_dict = ast.literal_eval(f.read())
    # Parse the text as JSON
    print(ai_response_dict)

    # print("AI response received:", ai_response_dict)
    return ai_response_dict, template_indicazioni, name, patient_id, analysis_type, committent, base_condition_filter, other_conditions_filter
    # fill_template_from_dict(template_indicazioni, ai_response_dict, './ARCHIVIO/{0}_{1}_{2}_indicazioni.docx'.format(name, patient_id, analysis_type), committent, analysis_type)

    # final_docx = './ARCHIVIO/{0}_{1}_{2}_result.docx'.format(name, patient_id, analysis_type)
    # merge_docx(['./ARCHIVIO/{0}_{1}_{2}_genetics.docx'.format(name, patient_id, analysis_type), './ARCHIVIO/{0}_{1}_{2}_indicazioni.docx'.format(name, patient_id, analysis_type)], final_docx)

    
    if button == "Invia ad Astrolabio":
        copyfile(final_pdf, os.path.join('./pseudoAstrolabio', os.path.basename(final_pdf)))


