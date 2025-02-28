import json
from docx import Document
import ast
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

def find_key(data, target_key):
    """Recursively search for the target_key in the JSON structure and return its value."""
    if isinstance(data, dict):
        for key, value in data.items():
            if key == target_key:
                return value
            elif isinstance(value, (dict, list)):
                found = find_key(value, target_key)
                if found is not None:
                    return found
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, (dict, list)):
                found = find_key(item, target_key)
                if found is not None:
                    return found
    return None

def paragraph_format_run(cell):
    """Helper function to create a paragraph, format and run in a cell."""
    cell_p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    cell_p_format = cell_p.paragraph_format
    cell_p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_r = cell_p.add_run()
    return cell_p, cell_p_format, cell_r


def clear_existing_header(header):
    """Clears all content from the header."""
    for paragraph in header.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._element = None

def add_image_to_header(doc, analysis_type, committent):
    for section in doc.sections:
        header = section.header

        # Clear existing header content to avoid duplication
        clear_existing_header(header)

        # Create a new table in the header
        htable = header.add_table(1, 2, Inches(7.6))
        htable.alignment = WD_TABLE_ALIGNMENT.LEFT
        htab_cells = htable.rows[0].cells

        # Left cell
        ht0 = htab_cells[0]
        ht0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell_p, cell_f, cell_r = paragraph_format_run(ht0)

        # Get the current directory of the script
        current_dir = os.path.dirname(os.path.abspath(__file__))

        # Construct the path to the image in the static folder
        static_path = os.path.join(current_dir, '..', 'static')

        # Normalize the path (resolves the ".." and any other relative path elements)
        static_path = os.path.normpath(static_path)
        # Add the appropriate image based on the 'committent'
        if committent == 'Genessere':
            cell_r.add_picture(f"{static_path}/header_logos/intestazione_genessere.png", width=Inches(2.55))
        else:
            cell_r.add_picture(f"{static_path}/header_logos/logo_{analysis_type}.png", width=Inches(1.19))

        # Right cell
        ht1 = htab_cells[1]
        ht1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell_p, cell_f, cell_r = paragraph_format_run(ht1)
        cell_r.add_picture(f"{static_path}/header_logos/intestazione_altamedica.png", width=Inches(3.33))
        cell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set autofit for the table
        htable.allow_autofit = True

def fill_template_from_dict(template_path, dict_data, output_path, committent, analysis_type):
    # Load the DOCX template
    doc = Document(template_path)
    
    integratori_data = find_key(dict_data, 'Integratori')

    # Handle 'Integratori' as either a dictionary or a list of dictionaries
    # integratori_data = dict_data.get('raccomandazioni', {}).get('Integratori', {})
    # print(integratori_data)
    if isinstance(integratori_data, dict):
        integratori_text = '\n\n'.join([f"Tipo: {v.get('tipo', '')}\nDosaggio: {v.get('dosaggio', '')}\nMotivazione: {v.get('motivazione', '')}" for k, v in integratori_data.items()]) #[f"{k} ({v.get('marca', '')}, {v.get('dosaggio', '')}): {v.get('motivazione', '')}" for k, v in integratori_data.items()])
    elif isinstance(integratori_data, list):
        integratori_text = '\n\n'.join('\n'.join([f"{key.capitalize()}: {value}" for key, value in integratore.items()]) for integratore in integratori_data)
    else:
        integratori_text = ''

    print("CONDIZIONI LISTA")
    updated_condizioni = []
    for condizione in dict_data.get('condizioni', []).split(', ') if isinstance(dict_data.get('condizioni', []), str) else dict_data.get('condizioni', []):
        if not any(keyword in condizione for keyword in ["Predisposizione", "predisposizione", "Intolleranza"]):
            updated_condizioni.append(f"Paziente riferisce: {condizione}")
        else:
            updated_condizioni.append(condizione)
    dict_data['condizioni'] = updated_condizioni
    
    print('\n'.join(dict_data.get('condizioni', [])))
    # Mapping placeholders to corresponding JSON keys, using empty strings for missing data
    placeholders = {
        "<condizioni>": '\n'.join([condizione.capitalize() for condizione in dict_data.get('condizioni', [])]),
        "<diagnosi>": dict_data.get('Diagnosi', ''),
        "<pesce_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('PESCE', {}).get('items', ''),
        "<pesce_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('PESCE', {}).get('items', ''),
        "<pesce_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('PESCE', {}).get('items', ''),
        "<carne_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('CARNE', {}).get('items', ''),
        "<carne_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('CARNE', {}).get('items', ''),
        "<carne_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('CARNE', {}).get('items', ''),
        "<veg_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('ALIMENTI DI ORIGINE VEGETALE', {}).get('items', ''),
        "<veg_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('ALIMENTI DI ORIGINE VEGETALE', {}).get('items', ''),
        "<veg_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('ALIMENTI DI ORIGINE VEGETALE', {}).get('items', ''),
        "<latt_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('LATTICINI', {}).get('items', ''),
        "<latt_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('LATTICINI', {}).get('items', ''),
        "<latt_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('LATTICINI', {}).get('items', ''),
        "<legumi_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('LEGUMI', {}).get('items', ''),
        "<legumi_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('LEGUMI', {}).get('items', ''),
        "<legumi_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('LEGUMI', {}).get('items', ''),
        "<uova_consigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Consigliati', {}).get('UOVA', {}).get('items', ''),
        "<uova_tollerati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Tollerati', {}).get('UOVA', {}).get('items', ''),
        "<uova_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Proteine', {}).get('Sconsigliati', {}).get('UOVA', {}).get('items', ''),
        "<cereali_consigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Consigliati', {}).get('CEREALI', {}).get('items', ''),
        "<cereali_tollerati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Tollerati', {}).get('CEREALI', {}).get('items', ''),
        "<cereali_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Sconsigliati', {}).get('CEREALI', {}).get('items', ''),
        "<pseudoc_consigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Consigliati', {}).get('PSEUDO-CEREALI', {}).get('items', ''),
        "<pseudoc_tollerati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Tollerati', {}).get('PSEUDO-CEREALI', {}).get('items', ''),
        "<pseudoc_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Sconsigliati', {}).get('PSEUDO-CEREALI', {}).get('items', ''),
        "<frutta_consigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Consigliati', {}).get('FRUTTA', {}).get('items', ''),
        "<frutta_tollerati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Tollerati', {}).get('FRUTTA', {}).get('items', ''),
        "<frutta_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Carboidrati', {}).get('Sconsigliati', {}).get('FRUTTA', {}).get('items', ''),
        "<grassi_insat_consigliati>": dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Consigliati', {}).get('GRASSI MONOINSATURI E POLINSATURI', {}).get('items', ''),
        "<grassi_insat_tollerati>": dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Tollerati', {}).get('GRASSI MONOINSATURI E POLINSATURI', {}).get('items', ''),
        "<grassi_insat_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Sconsigliati', {}).get('GRASSI MONOINSATURI E POLINSATURI', {}).get('items', ''),
        "<grassi_sat_consigliati>":  dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Consigliati', {}).get('GRASSI SATURI', {}).get('items', ''),
        "<grassi_sat_tollerati>": dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Tollerati', {}).get('GRASSI SATURI', {}).get('items', ''),
        "<grassi_sat_sconsigliati>": dict_data.get('raccomandazioni', {}).get('Lipidi', {}).get('Sconsigliati', {}).get('GRASSI SATURI', {}).get('items', ''),
        "<verdure>": dict_data.get('raccomandazioni', {}).get('Verdure', ''),
        "<integratori>": integratori_text,
    }
    
    add_image_to_header(doc, analysis_type, committent)

    # Replace placeholders in the document
    def replace_text_in_paragraph(paragraph, placeholders):
        """Replaces the text in a paragraph based on the given placeholders."""
        # print("\n\n###################### TEXT REPLACEMENT ######################\n\n")
        for placeholder, replacement in placeholders.items():
            # print("\n\n###################### BEFORE REPLACEMENT ######################")
            # print(replacement)
            # replacement = replacement.capitalize()
            replacement = replacement.replace('<NEWLINE>', '\n')
            replacement = replacement.replace('<newline>', '\n')
            # print("\n\n###################### REPLACEMENT DONE ######################")
            # print(replacement)
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    def replace_text_in_table(table, placeholders):
        """Replaces the text in a table's cells based on the given placeholders."""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, placeholders)
    
    # Replace text in the paragraphs (outside tables)
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)
    
    # Replace text in all tables
    for table in doc.tables:
        replace_text_in_table(table, placeholders)
    
    # Save the modified document
    doc.save(output_path)
    print(f"Document saved as {output_path}")

# Example usage with json.loads()
# json_str = '''{
#     "id_paziente": "T2D_Lieve+Cardio_NonEvidente+Peso_NonEvidente+FerroBasso_True",
#     "condizioni": ["Diabete di tipo 2 (Lieve)", "Predisposizione malattie cardiovascolari (NonEvidente)", "Predisposizione all'aumento di peso (NonEvidente)", "PREDISPOSIZIONE GENETICA ALLA CARENZA DI FERRO"],
#     ...
# }'''

# json_data = json.loads(json_str)
# with open('/home/mauro/Desktop/butta/response.txt', 'r', encoding='utf-8') as file:
#         file_content = file.read()

#         json_data = ast.literal_eval(file_content)
# fill_template_from_dict('/home/mauro/Desktop/GoogleDrive/Work/Altamedica/Genefood/static/Indicazioni_alimentari.docx', json_data, '/home/mauro/Desktop/butta/Filled_Indicazioni_alimentari.docx', committent='Altamedica', analysis_type='Vita')
# print(json_data.get('Verdure', ''))