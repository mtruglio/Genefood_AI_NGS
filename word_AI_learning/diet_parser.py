import docx
import pandas as pd
import os
import json
from pathlib import Path
import sys
import re

# Translation dictionaries
condition_translation = {
    "T2D": "Diabete di tipo 2",
    "Cardio": "Predisposizione malattie cardiovascolari",
    "Peso": "Predisposizione all'aumento di peso",
    "Latt": "Intolleranza al lattosio",
    "Glut": "Intolleranza al glutine",
    "Vegetariano": "Dieta vegetariana",
    "Pescetariano": "Dieta pescetariana",
    "Vegano": "Dieta vegana"
}

risk_translation = {
    "Non Evidente": "Non evidente",
    "Lieve": "Lieve",
    "Medio": "Medio",
    "Alto": "Alto"
}

def clean_table(table):
    cleaned_table = []
    for row in table:
        if all(cell.strip() == '' for cell in row):
            continue
        cleaned_row = []
        for i in range(len(row)):
            current_cell = row[i].strip()
            if i < len(row) - 1:
                next_cell = row[i + 1].strip()
                if current_cell == '' and next_cell == '':
                    cleaned_row.append('')
                elif current_cell:
                    cleaned_row.append(current_cell)
            else:
                if current_cell:
                    cleaned_row.append(current_cell)
        if cleaned_row:
            cleaned_table.append(cleaned_row)
    return cleaned_table

def extract_and_clean_tables(doc_path):
    doc = docx.Document(doc_path)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    cleaned_tables = [clean_table(table) for table in tables[1:4]]
    return cleaned_tables

def parse_table_with_subcategories(table):
    parsed_data = {"Consigliati": {}, "Tollerati": {}, "Sconsigliati": {}}

    for i, row in enumerate(table):
        if i == 0:
            continue
        sub_category = row[0].split('\n')[0]
        dosage = ' '.join(row[0].split('\n')[1:])
        consigliati = row[1] if len(row) > 1 else ""
        tollerati = row[2] if len(row) > 2 else ""
        sconsigliati = row[3] if len(row) > 3 else ""

        latte_text = ""
        if "Latte" in dosage:
            if "C:" not in dosage:
                sys.exit("ERROR: Latte found but no C: found")
            dosage = dosage.replace("Yogurt", "yogurt")
            latte_text = dosage.split('Latte e yogurt:')[1].split('C:')[0].strip()

        def extract_relevant_dosage(dosage, category):
            parts = dosage.split()
            relevant_dosage = []
            capture = False
            for part in parts:
                if part.startswith(category + ':'):
                    capture = True
                    relevant_dosage.append(part)
                elif capture and (part.startswith('C:') or part.startswith('T:') or part.startswith('S:')):
                    break
                elif capture:
                    relevant_dosage.append(part)
            return ' '.join(relevant_dosage)

        parsed_data["Consigliati"][sub_category] = {
            "dosage": f"Latte e yogurt: {latte_text}. " + extract_relevant_dosage(dosage, 'C').strip() if latte_text else extract_relevant_dosage(dosage, 'C').strip(),
            "items": ", ".join([item.strip() for item in consigliati.split(',') if item.strip()])
        }
        parsed_data["Tollerati"][sub_category] = {
            "dosage": f"Latte e yogurt: {latte_text}. " + extract_relevant_dosage(dosage, 'T').strip() if latte_text else extract_relevant_dosage(dosage, 'T').strip(),
            "items": ", ".join([item.strip() for item in tollerati.split(',') if item.strip()])
        }
        parsed_data["Sconsigliati"][sub_category] = {
            "dosage": f"Latte e yogurt: {latte_text}. " + extract_relevant_dosage(dosage, 'S').strip() if latte_text else extract_relevant_dosage(dosage, 'S').strip(),
            "items": ", ".join([item.strip() for item in sconsigliati.split(',') if item.strip()])
        }
    return parsed_data

def extract_verdure_section(doc_path):
    doc = docx.Document(doc_path)
    verdure_text = ""
    capture = False

    for para in doc.paragraphs:
        if "Verdure consigliate" in para.text:
            capture = True
        if capture:
            verdure_text += para.text + " "

    return verdure_text.strip()

def extract_diagnosi_section(doc_path):
    doc = docx.Document(doc_path)
    diagnosi_text = ""
    capture = False

    for para in doc.paragraphs:
        if "Il tuo profilo" in para.text:
            capture = True
        if capture:
            diagnosi_text += para.text + " "
            next_element = para._element.getnext()
            if next_element is not None and next_element.tag.endswith('tbl'):
                break
    
    return diagnosi_text.strip()

def translate_conditions(conditions):
    translated_conditions = []
    for condition in conditions:
        parts = condition.split('_')
        condition_name = parts[0]
        risk_level = parts[1] if len(parts) > 1 else ""
        translated_condition = condition_translation.get(condition_name, condition_name)
        translated_risk = risk_translation.get(risk_level, risk_level)
        if translated_risk:
            translated_conditions.append(f"{translated_condition} ({translated_risk})")
        else:
            translated_conditions.append(translated_condition)
    return translated_conditions

def process_documents_with_conditions(folder_path):
    doc_paths = Path(folder_path).glob("*[+]*.docx")
    data = []
    for doc_path in doc_paths:
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        conditions = base_name.split('+')
        translated_conditions = translate_conditions(conditions)
        cleaned_tables = extract_and_clean_tables(doc_path)
        verdure_text = extract_verdure_section(doc_path)
        diagnosi_text = extract_diagnosi_section(doc_path)

        recommendations = {}
        categories = ["Proteine", "Carboidrati", "Lipidi", "Verdure"]
        for i, table in enumerate(cleaned_tables):
            recommendations[categories[i]] = parse_table_with_subcategories(table)

        recommendations["Verdure"] = verdure_text

        patient_data = {
            "id_paziente": base_name,
            "condizioni": translated_conditions,
            "raccomandazioni": recommendations,
            "Diagnosi": diagnosi_text
        }
        data.append(patient_data)
    return data

def extract_diagnosi_and_raccomandazioni(doc_path):
    doc = docx.Document(doc_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    diagnosi_start = text.find("Il tuo profilo genetico ha evidenziato")
    diagnosi_end = text.find("Consigli", diagnosi_start)
    diagnosi_text = text[diagnosi_start:diagnosi_end].strip()
    
    raccomandazioni_start = text.find("Consigli alimentari in base al tuo profilo genetico:")
    raccomandazioni_end = text.find("Bibliografia", raccomandazioni_start)
    raccomandazioni_text = text[raccomandazioni_start:raccomandazioni_end].strip()
    
    return diagnosi_text, raccomandazioni_text

def process_documents_without_conditions(folder_path):
    doc_paths = Path(folder_path).glob("*.docx")
    data = []
    for doc_path in doc_paths:
        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        if "+" in base_name:
            continue
        diagnosi_text, raccomandazioni_text = extract_diagnosi_and_raccomandazioni(doc_path)
        
        patient_data = {
            "id_paziente": base_name,
            "condizioni": base_name.replace("_", " "),
            "raccomandazioni": raccomandazioni_text,
            "Diagnosi": diagnosi_text
        }
        data.append(patient_data)
    return data

def process_documents_to_json(folder_path):
    data = process_documents_with_conditions(folder_path)
    data += process_documents_without_conditions(folder_path)

    with open('consolidated_data_italian_with_subcategories_junior_more.json', 'w') as json_file:
        json.dump(data, json_file, indent=4, ensure_ascii=False)
    
    with open('consolidated_data_italian_with_subcategories_junior_more.json', 'r') as file:
        filedata = file.read()

    filedata = filedata.replace('"."', '""').replace('\\n', ' ').replace('\\t', ' ')

    def replace_spaces_in_quotes(match):
        return re.sub(' +', ' ', match.group(0))

    pattern = r'"(?:\\.|[^"\\])*"'

    filedata = re.sub(pattern, replace_spaces_in_quotes, filedata)
    with open('consolidated_data_italian_with_subcategories_junior_more.json', 'w') as file:
        file.write(filedata)

if __name__ == "__main__":
    folder_path = sys.argv[1]
    process_documents_to_json(folder_path)
